import json
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOC_DIR = ROOT / "docs" / "e2e_testing"
SCREEN_DIR = DOC_DIR / "screenshots"
DATA = json.loads((DOC_DIR / "guide_data.json").read_text(encoding="utf-8"))
OUT_DOCX = DOC_DIR / "EmoScreen_End_to_End_Testing_Guide.docx"
LIVE_BASE_URL = "https://emo.cpdinclinic.co.in"


def live_url(path):
    return f"{LIVE_BASE_URL}{path}"

BLUE = RGBColor(47, 62, 158)
TEAL = RGBColor(42, 167, 161)
DARK = RGBColor(31, 45, 61)
MUTED = RGBColor(107, 124, 147)
LIGHT_TEAL = "EAF7F7"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F6FAFB"
BORDER = "D8E7EB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc, text="", *, style=None, size=11, color=DARK, bold=False, italic=False, align=None, after=6):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_title(doc, title, subtitle):
    add_paragraph(doc, "EmoScreen", size=12, color=TEAL, bold=True, after=4)
    add_paragraph(doc, title, size=28, color=BLUE, bold=True, after=6)
    add_paragraph(doc, subtitle, size=14, color=MUTED, after=14)
    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1800, 7560])
    rows = [
        ("Audience", "Clinic teams, pediatricians, caregivers, support users, and UAT testers"),
        ("Scope", "Admin doctor recruitment, doctor Google-gated sharing, paid parent flow, free multilingual parent flow, self-screening, reports, and audit visibility"),
        ("Environment", "Written as a live-system walkthrough. Screenshots were captured from the local UAT system."),
        ("Guide date", date.today().strftime("%B %d, %Y")),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
    doc.add_paragraph()
    add_callout(
        doc,
        "How to read this guide",
        "Follow the screens in order. On each screen, perform the listed action and confirm the expected result before moving ahead. For live production, use the production domain and authorized accounts. For local UAT, the payment screen uses the dummy payment button instead of live Razorpay.",
    )


def add_callout(doc, heading, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(heading)
    set_run(r, size=11, color=TEAL, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=DARK)
    doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run(r, size=16, color=BLUE, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run(r, size=13, color=BLUE, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run(r, size=12, color=RGBColor(31, 77, 120), bold=True)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run(r, size=10.5, color=DARK)


def add_step_table(doc, what_to_do, expected):
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [1900, 7460])
    rows = [("What to do", what_to_do), ("Expected result", expected)]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run(r, size=10.5, color=BLUE, bold=True)
        p2 = row.cells[1].paragraphs[0]
        r2 = p2.add_run(value)
        set_run(r2, size=10.5, color=DARK)
    doc.add_paragraph()


def add_screen(doc, title, image_name, what_to_do, expected, note=None):
    add_h3(doc, title)
    if note:
        add_paragraph(doc, note, size=10.5, color=MUTED, italic=True, after=6)
    add_step_table(doc, what_to_do, expected)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(SCREEN_DIR / image_name), width=Inches(5.75))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(12)
    r = caption.add_run(f"Screenshot: {title}")
    set_run(r, size=9, color=MUTED, italic=True)


def add_checklist_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1700, 3160, 2500, 2000])
    headers = ["Area", "Tester action", "Pass condition", "Result"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        r = cell.paragraphs[0].add_run(header)
        set_run(r, size=10, color=BLUE, bold=True)
    rows = [
        ("Admin recruitment", "Upload or register doctors from admin tools", "Doctor receives clinic link and how-to guide", "Pass / Fail"),
        ("Doctor access", "Doctor opens clinic link and signs in with registered Google email", "Doctor can see provider banner and sharing form", "Pass / Fail"),
        ("Paid form sharing", "Select paid form and send link", "Parent receives/opens payment and assessment link", "Pass / Fail"),
        ("Free form sharing", "Select behavioral form and language", "Parent verifies phone and reaches language/form flow", "Pass / Fail"),
        ("Payment", "Complete payment", "Payment status becomes completed before assessment opens", "Pass / Fail"),
        ("One-time use", "Reopen a completed paid/free form link", "Existing report opens; answers cannot be submitted again", "Pass / Fail"),
        ("Reports", "Confirm paid and free result/report screens", "Patient receives patient report; doctor receives doctor and patient reports where applicable", "Pass / Fail"),
        ("Audit trail", "Open support workflow dashboard and detail page", "Status, payment, report, event history, and delivery attempts are visible to support", "Pass / Fail"),
    ]
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            r = cell.paragraphs[0].add_run(value)
            set_run(r, size=9.5, color=DARK)
    set_table_geometry(table, [1700, 3160, 2500, 2000])
    doc.add_paragraph()


def add_language_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1800, 2300, 3360, 1900])
    headers = ["Language", "Screening link", "Tester action", "Expected result"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        r = cell.paragraphs[0].add_run(header)
        set_run(r, size=10, color=BLUE, bold=True)
    rows = [
        ("English", "Free form", "Select English and submit all questions", "Result page opens"),
        ("Hindi", "Free form", "Select Hindi and submit all questions", "Result page opens"),
        ("Tamil", "Free form", "Select Tamil and submit all questions", "Result page opens"),
        ("Telugu", "Free form", "Select Telugu and submit all questions", "Result page opens"),
        ("Malayalam", "Free form", "Select Malayalam and submit all questions", "Result page opens"),
        ("Marathi", "Free form", "Select Marathi and submit all questions", "Result page opens"),
        ("Bengali", "Free form", "Select Bengali and submit all questions", "Result page opens"),
        ("Kannada", "Free form", "Select Kannada and submit all questions", "Result page opens"),
    ]
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            r = cell.paragraphs[0].add_run(value)
            set_run(r, size=9.5, color=DARK)
    set_table_geometry(table, [1800, 2300, 3360, 1900])
    doc.add_paragraph()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    header = section.header.paragraphs[0]
    header.text = "EmoScreen End-to-End Testing Guide"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("User-oriented UAT walkthrough"), size=9, color=MUTED)
    return doc


def build_doc():
    doc = setup_document()
    add_title(
        doc,
        "End-to-End Testing Guide",
        "A user-oriented walkthrough for recruitment, doctor sharing, parent completion, reports, and support tracking",
    )

    add_h1(doc, "1. Test Overview")
    add_paragraph(
        doc,
        "This guide walks a tester through EmoScreen as separate real users experience it: public visitor, admin, recruited doctor, parent, and support user.",
    )
    add_bullets(
        doc,
        [
            "Public visitors enter from the normal EmoScreen home page.",
            "Admins enter separately through Django admin and recruitment tools.",
            "Doctors do not use admin login for their clinical workflow; they open their clinic link and sign in with the registered Google email.",
            "Parents follow only the paid, free, or self-screening journey they received.",
            "Support users verify delivery, payment, submission, and report status in the workflow audit dashboard.",
        ],
    )
    add_callout(
        doc,
        "Role separation rule",
        "Do not test doctor sharing from the Django admin dashboard. Admin is for recruitment/support. Doctor sharing starts from the doctor clinic link received after recruitment.",
    )
    add_callout(
        doc,
        "Payment note for UAT",
        "The live system should use Razorpay for paid forms. In this local UAT guide, the screen uses Complete Dummy Payment so testers can validate the full journey without moving real money.",
    )

    add_h2(doc, "Entry points")
    doctor_live_clinic_url = live_url(f"/admin/bulk-upload/clinic/{DATA['doctor_code']}/")
    entry = doc.add_table(rows=4, cols=3)
    set_table_geometry(entry, [2200, 3300, 3860])
    headers = ["User", "Live-style entry", "Purpose"]
    for cell, header in zip(entry.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        r = cell.paragraphs[0].add_run(header)
        set_run(r, size=10, color=BLUE, bold=True)
    rows = [
        ("Normal user / parent", live_url("/"), "Public landing and patient start points"),
        ("Admin / support", live_url("/admin/"), "Django admin, doctor recruitment, reports, and staff-only tools"),
        ("Doctor", doctor_live_clinic_url, "Google-gated clinic console sent to the recruited doctor"),
    ]
    for row, values in zip(entry.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            r = cell.paragraphs[0].add_run(value)
            set_run(r, size=9.5, color=DARK)
    doc.add_paragraph()

    add_h2(doc, "Sample UAT details")
    sample = doc.add_table(rows=5, cols=2)
    set_table_geometry(sample, [2200, 7160])
    rows = [
        ("Doctor workspace", f"Dr Codex QA, doctor code {DATA['doctor_code']}"),
        ("Paid sample patient", "Aarav Guide Test"),
        ("Free sample patient", "Meera Free Guide"),
        ("Paid workflow case", DATA["paid_case_code"]),
        ("Free workflow case", DATA["free_case_code"]),
    ]
    for row, (label, value) in zip(sample.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    add_h1(doc, "2. Public Entry")
    add_screen(
        doc,
        "Public User Entry",
        "02-public-entry.png",
        "Open the normal EmoScreen entry page. Use this only for public registration, public start, or self-screening routes.",
        "The public page is reachable without staff/admin credentials and does not replace the staff admin entry.",
    )

    add_h1(doc, "3. Admin Recruitment Workflow")
    add_screen(
        doc,
        "Admin Sign-In",
        "01-admin-sign-in.png",
        "Admin or support user signs in from the admin entry point using a staff account.",
        "The admin session is opened only for recruitment, reports, and support operations.",
    )
    add_screen(
        doc,
        "Bulk Doctor Recruitment",
        "03-admin-bulk-doctor-recruitment.png",
        "Open the bulk doctor upload screen, choose a CSV with doctor details, and submit it.",
        "Valid doctors are created, duplicates are skipped, failures are shown, and a result CSV is available.",
    )
    add_callout(
        doc,
        "Doctor onboarding message",
        f"After recruitment the doctor receives WhatsApp/email copy like: Hello Dr Name. Thank you for registering for the Emotional & Behavioural Screening Tool - an initiative supported by SAPA. Your personalized screening tool is now ready. Clinic Link: {doctor_live_clinic_url} How to Use Guide: {DATA['how_to_guide_url']}.",
    )

    add_h1(doc, "4. Doctor Workflow")
    add_callout(
        doc,
        "Doctor sign-in rule",
        "In production the doctor opens the clinic link from the onboarding message and signs in with Google using the same email that was registered. If another Google account is used, access should be blocked.",
    )
    add_screen(
        doc,
        "Clinic Console",
        "04-doctor-clinic-console.png",
        "Doctor confirms the provider banner shows the correct clinic details. Use Send Form to Parent to share free or paid forms.",
        "The doctor sees only the clinic sharing workspace and can choose a parent WhatsApp number, form type, language, and QR/direct sharing options.",
    )
    add_screen(
        doc,
        "Doctor Shares a Paid Form",
        "05-doctor-share-paid-form.png",
        "Choose a paid age-band form, enter the parent's WhatsApp number, enter the patient name, choose the price, and send via WhatsApp.",
        "WhatsApp opens with a prepared message containing the paid parent link.",
    )
    add_screen(
        doc,
        "Doctor Shares a Free Form",
        "06-doctor-share-free-form.png",
        "Choose the Behavioral and Emotional Red Flags form, select the parent's preferred language, and send via WhatsApp.",
        "WhatsApp opens with a prepared message containing the verification link in the selected language.",
    )

    add_h1(doc, "5. Paid Parent Flow")
    add_screen(
        doc,
        "Parent Opens Paid Order",
        "07-parent-paid-order-entry.png",
        "Parent opens the paid link and confirms or enters their email address before continuing.",
        "The order page shows the prescribed form and payable amount.",
    )
    add_screen(
        doc,
        "Payment Screen",
        "08-payment-screen.png",
        "For local UAT, click Complete Dummy Payment. In production, the parent should complete the Razorpay checkout.",
        "Payment is marked completed and the parent is allowed to proceed to the assessment.",
    )
    add_screen(
        doc,
        "Paid Assessment Form",
        "09-paid-assessment-form.png",
        "Parent fills child details, confirms consent, answers all questions, and clicks Review.",
        "The Review screen opens only after required fields and questions are completed.",
    )
    add_screen(
        doc,
        "Review Answers",
        "10-paid-review-answers.png",
        "Parent reviews answers. If correct, click Submit Final. If incorrect, click Edit Answers.",
        "Final submission triggers report generation.",
    )
    add_screen(
        doc,
        "Paid Report Screen",
        "11-paid-report-screen.png",
        "Confirm that the submission was accepted and the patient report screen is shown.",
        "The patient can download only the patient report. The doctor receives the doctor report and patient report by email. Reopening the paid link returns to this report screen instead of allowing another submission.",
        note="In local UAT, email delivery may show FAILED when provider credentials are absent. In production, support should expect provider delivery status.",
    )

    add_h1(doc, "6. Free Parent Flow")
    add_screen(
        doc,
        "Phone Verification",
        "12-free-phone-verification.png",
        "Parent enters the same WhatsApp number that received the doctor message.",
        "The parent is verified and moved to language selection.",
    )
    add_screen(
        doc,
        "Language Selection",
        "13-language-selection.png",
        "Parent selects the language they are most comfortable using.",
        "The selected language opens the screening form.",
    )
    add_screen(
        doc,
        "Free Screening Form",
        "14-free-screening-form.png",
        "Parent enters patient details, answers every question, and submits the form.",
        "The result page opens after all required fields are complete.",
    )
    add_screen(
        doc,
        "Free Result Screen",
        "15-free-result-screen.png",
        "Parent reviews the result summary and keeps the report code for future reference.",
        "The result clearly shows whether warning signs were identified and shows next-step actions where applicable. Reopening the completed verified link returns to the existing report instead of accepting another submission.",
    )

    add_h2(doc, "Language coverage checklist")
    add_paragraph(
        doc,
        "Run the same free screening submission once in each supported language. The tester does not need to understand every translated question; the test is passed when the language-specific form loads, accepts answers, and opens the result screen.",
    )
    add_language_table(doc)

    add_h1(doc, "7. Self-Screening Flow")
    add_screen(
        doc,
        "Self-Screen Start",
        "16-self-screen-start.png",
        "A parent without a doctor link starts from Self Screening and enters their WhatsApp number.",
        "The system routes the parent into the same guided screening experience without assigning a clinic doctor.",
    )

    add_h1(doc, "8. Support and Audit Visibility")
    add_screen(
        doc,
        "Workflow Audit Dashboard",
        "17-workflow-audit-dashboard.png",
        "Support opens the workflow dashboard and searches/filters by doctor, patient, form status, payment status, report status, or date range.",
        "Support can identify how many workflows are pending, completed, failed, paid, or submitted.",
    )
    add_screen(
        doc,
        "Workflow Detail Page",
        "18-workflow-audit-detail.png",
        "Open a workflow case to inspect doctor, patient, payment, report, event timeline, and delivery attempts.",
        "Support can answer where the workflow currently stands and where it failed if any stage did not complete.",
    )

    add_h1(doc, "9. Final Pass Checklist")
    add_paragraph(doc, "Use this checklist at the end of a full live or UAT pass.")
    add_checklist_table(doc)

    add_h1(doc, "10. Completion Criteria")
    add_bullets(
        doc,
        [
            "Normal public entry and admin entry are separate and both remain reachable.",
            "Admin can recruit doctors and the doctor receives the clinic link plus how-to guide.",
            "Doctor clinic workflow starts from the onboarding clinic link and requires the registered Google email.",
            "Doctor can share both paid and free links without confusion.",
            "Parent can complete payment before accessing a paid assessment.",
            "Parent can complete the paid form once and reach the patient-facing paid report screen.",
            "Paid report delivery sends one patient report to the patient and both patient and doctor reports to the doctor.",
            "Parent can complete the free form once in all supported languages.",
            "Self-screening starts without a doctor-specific link.",
            "Support can trace each workflow from creation through completion using the audit dashboard.",
            "Any failure state shows the stage, reason, and enough context for support to retry or escalate.",
        ],
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
